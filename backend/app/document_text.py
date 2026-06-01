from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader


TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".json"}
PDF_CONTENT_TYPES = {"application/pdf", "application/x-pdf"}
DOCX_CONTENT_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
}


class DocumentTextError(ValueError):
    pass


def extract_document_text(filename: str | None, content_type: str | None, data: bytes) -> str:
    suffix = Path(filename or "").suffix.lower()
    normalized_content_type = (content_type or "").lower().split(";")[0].strip()

    if suffix == ".pdf" or normalized_content_type in PDF_CONTENT_TYPES:
        return extract_pdf_text(data)
    if suffix == ".docx" or normalized_content_type in DOCX_CONTENT_TYPES:
        return extract_docx_text(data)
    if suffix in TEXT_EXTENSIONS or normalized_content_type.startswith("text/") or not suffix:
        return decode_text_bytes(data)

    raise DocumentTextError("Unsupported file type. Please upload TXT, Markdown, CSV, PDF, or DOCX.")


def decode_text_bytes(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-16", "gb18030"):
        try:
            return data.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore").strip()


def extract_pdf_text(data: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(data))
    except Exception as exc:
        raise DocumentTextError("PDF file could not be opened.") from exc

    text = "\n".join(page.extract_text() or "" for page in reader.pages).strip()
    if not text:
        raise DocumentTextError("PDF text could not be extracted. Scanned PDFs need OCR before upload.")
    return text


def extract_docx_text(data: bytes) -> str:
    try:
        document = Document(BytesIO(data))
    except Exception as exc:
        raise DocumentTextError("DOCX file could not be opened.") from exc

    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    if not text:
        raise DocumentTextError("DOCX file does not contain readable text.")
    return text
